import os
import sys
import docx
import numpy as np
from sentence_transformers import SentenceTransformer

def extract_jd_text(file_path: str) -> str:
    """
    Extracts text from the job description docx.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Job description file not found: {file_path}")
        
    doc = docx.Document(file_path)
    text_parts = []
    
    # Extract paragraphs
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            text_parts.append(t)
            
    # Extract table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t and t not in text_parts:
                    text_parts.append(t)
                    
    return "\n".join(text_parts)

def precompute_jd_embedding():
    workspace_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    jd_path = os.path.join(workspace_dir, "data", "job_description.docx")
    out_path = os.path.join(workspace_dir, "artifacts", "jd_embedding.npy")
    
    print(f"Reading JD from {jd_path}...")
    jd_text = extract_jd_text(jd_path)
    print(f"Extracted JD text length: {len(jd_text)} characters")
    
    print("Loading SentenceTransformer model 'all-MiniLM-L6-v2'...")
    # This will download the model locally on first run and cache it
    model = SentenceTransformer('all-MiniLM-L6-v2')
    
    print("Computing JD embedding...")
    jd_embedding = model.encode(jd_text, show_progress_bar=True)
    
    print(f"JD embedding shape: {jd_embedding.shape}")
    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    np.save(out_path, jd_embedding)
    print(f"Successfully saved JD embedding to {out_path}")

if __name__ == "__main__":
    precompute_jd_embedding()
